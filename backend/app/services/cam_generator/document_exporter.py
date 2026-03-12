"""Document exporter for converting CAM documents to Word and PDF formats."""

import os
from datetime import datetime
from typing import Optional
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib import colors
from reportlab.pdfgen import canvas

from app.models.cam import CAMDocument


class DocumentExporter:
    """Exports CAM documents to Word (.docx) and PDF formats."""

    def __init__(self, template_dir: Optional[str] = None):
        """Initialize DocumentExporter.
        
        Args:
            template_dir: Optional directory containing document templates
        """
        self.template_dir = template_dir or os.path.join(
            os.path.dirname(__file__), "templates"
        )
        self._ensure_template_dir()

    def _ensure_template_dir(self) -> None:
        """Ensure template directory exists."""
        Path(self.template_dir).mkdir(parents=True, exist_ok=True)

    def export_to_word(
        self, cam_document: CAMDocument, output_path: str
    ) -> None:
        """Export CAM document to Word (.docx) format.
        
        Args:
            cam_document: CAMDocument to export
            output_path: Path where the Word file will be saved
            
        Raises:
            IOError: If file cannot be written
            ValueError: If CAM document is invalid
        """
        if not cam_document.sections:
            raise ValueError("CAM document has no sections to export")

        # Create parent directories if needed
        output_dir = os.path.dirname(output_path)
        if output_dir:
            Path(output_dir).mkdir(parents=True, exist_ok=True)

        # Create a new Document
        doc = Document()

        # Add document title
        title_text = f"Credit Appraisal Memo - {self._sanitize_text(cam_document.company_name)}"
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add metadata
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(f"Application ID: ").bold = True
        metadata_para.add_run(self._sanitize_text(cam_document.application_id))
        
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(f"Generated Date: ").bold = True
        metadata_para.add_run(
            cam_document.generated_date.strftime("%d-%m-%Y %H:%M:%S")
        )
        
        metadata_para = doc.add_paragraph()
        metadata_para.add_run(f"Document Version: ").bold = True
        metadata_para.add_run(str(cam_document.version))

        doc.add_paragraph()  # Add spacing

        # Add sections
        section_order = [
            "executive_summary",
            "company_overview",
            "industry_analysis",
            "financial_analysis",
            "risk_assessment",
            "five_cs_summary",
            "final_recommendation",
            "explainability_notes",
            "audit_trail",
        ]

        for section_key in section_order:
            if section_key in cam_document.sections:
                section_content = cam_document.sections[section_key]
                self._add_section_to_word(doc, section_key, section_content)

        # Add footer with company branding
        self._add_footer_to_word(doc)

        # Save document
        try:
            doc.save(output_path)
        except IOError as e:
            raise IOError(f"Failed to save Word document to {output_path}: {str(e)}")

    def _add_section_to_word(
        self, doc: Document, section_key: str, content: str
    ) -> None:
        """Add a section to the Word document.
        
        Args:
            doc: python-docx Document object
            section_key: Key identifying the section
            content: Section content text
        """
        # Add section heading
        section_title = self._format_section_title(section_key)
        section_title = self._sanitize_text(section_title)
        doc.add_heading(section_title, level=1)

        # Parse and add content
        lines = content.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith("- "):
                # Bullet point
                sanitized_line = self._sanitize_text(line[2:])
                doc.add_paragraph(sanitized_line, style="List Bullet")
            elif line.endswith(":"):
                # Subheading
                sanitized_line = self._sanitize_text(line)
                doc.add_paragraph(sanitized_line, style="Heading 2")
            else:
                # Regular paragraph
                sanitized_line = self._sanitize_text(line)
                doc.add_paragraph(sanitized_line)

        doc.add_paragraph()  # Add spacing between sections

    def _format_section_title(self, section_key: str) -> str:
        """Format section key to readable title.
        
        Args:
            section_key: Section key (e.g., "executive_summary")
            
        Returns:
            Formatted title (e.g., "Executive Summary")
        """
        return " ".join(word.capitalize() for word in section_key.split("_"))

    def _sanitize_text(self, text: str) -> str:
        """Sanitize text to remove XML-incompatible characters.
        
        Args:
            text: Text to sanitize
            
        Returns:
            Sanitized text safe for XML/Word documents and ReportLab
        """
        if not text:
            return ""
        
        # Remove control characters (0x00-0x1F except tab, newline, carriage return)
        sanitized = ""
        for char in text:
            code = ord(char)
            # Allow tab (0x09), newline (0x0A), carriage return (0x0D)
            if code < 0x20 and code not in (0x09, 0x0A, 0x0D):
                # Replace control characters with space
                sanitized += " "
            else:
                sanitized += char
        
        # Escape XML special characters for ReportLab
        sanitized = sanitized.replace("&", "&amp;")
        sanitized = sanitized.replace("<", "&lt;")
        sanitized = sanitized.replace(">", "&gt;")
        
        return sanitized

    def _add_footer_to_word(self, doc: Document) -> None:
        """Add footer with company branding to Word document.
        
        Args:
            doc: python-docx Document object
        """
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Intelli-Credit AI Corporate Credit Decisioning Engine"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def export_to_pdf(
        self, cam_document: CAMDocument, output_path: str
    ) -> None:
        """Export CAM document to PDF format.
        
        Args:
            cam_document: CAMDocument to export
            output_path: Path where the PDF file will be saved
            
        Raises:
            IOError: If file cannot be written
            ValueError: If CAM document is invalid
        """
        if not cam_document.sections:
            raise ValueError("CAM document has no sections to export")

        # Create PDF document
        pdf_path = Path(output_path)
        pdf_path.parent.mkdir(parents=True, exist_ok=True)

        try:
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=0.75 * inch,
                leftMargin=0.75 * inch,
                topMargin=0.75 * inch,
                bottomMargin=0.75 * inch,
            )

            # Build story (content)
            story = []
            styles = getSampleStyleSheet()

            # Add title
            title_style = ParagraphStyle(
                "CustomTitle",
                parent=styles["Heading1"],
                fontSize=24,
                textColor=colors.HexColor("#1a1a1a"),
                spaceAfter=12,
                alignment=1,  # CENTER
            )
            title_text = f"Credit Appraisal Memo - {self._sanitize_text(cam_document.company_name)}"
            story.append(Paragraph(title_text, title_style))
            story.append(Spacer(1, 0.2 * inch))

            # Add metadata
            metadata_style = ParagraphStyle(
                "Metadata",
                parent=styles["Normal"],
                fontSize=10,
                textColor=colors.HexColor("#666666"),
            )
            story.append(
                Paragraph(
                    f"<b>Application ID:</b> {self._sanitize_text(cam_document.application_id)}",
                    metadata_style,
                )
            )
            story.append(
                Paragraph(
                    f"<b>Generated Date:</b> {cam_document.generated_date.strftime('%d-%m-%Y %H:%M:%S')}",
                    metadata_style,
                )
            )
            story.append(
                Paragraph(
                    f"<b>Document Version:</b> {cam_document.version}",
                    metadata_style,
                )
            )
            story.append(Spacer(1, 0.3 * inch))

            # Add sections
            section_order = [
                "executive_summary",
                "company_overview",
                "industry_analysis",
                "financial_analysis",
                "risk_assessment",
                "five_cs_summary",
                "final_recommendation",
                "explainability_notes",
                "audit_trail",
            ]

            for section_key in section_order:
                if section_key in cam_document.sections:
                    section_content = cam_document.sections[section_key]
                    self._add_section_to_pdf(story, section_key, section_content, styles)
                    story.append(PageBreak())

            # Add footer
            story.append(Spacer(1, 0.2 * inch))
            footer_style = ParagraphStyle(
                "Footer",
                parent=styles["Normal"],
                fontSize=9,
                textColor=colors.HexColor("#999999"),
                alignment=1,  # CENTER
            )
            story.append(
                Paragraph(
                    "Intelli-Credit AI Corporate Credit Decisioning Engine",
                    footer_style,
                )
            )

            # Build PDF
            doc.build(story)

        except IOError as e:
            raise IOError(f"Failed to save PDF document to {output_path}: {str(e)}")

    def _add_section_to_pdf(
        self,
        story: list,
        section_key: str,
        content: str,
        styles: dict,
    ) -> None:
        """Add a section to the PDF story.
        
        Args:
            story: ReportLab story list
            section_key: Key identifying the section
            content: Section content text
            styles: ReportLab styles dictionary
        """
        # Add section heading
        section_title = self._format_section_title(section_key)
        heading_style = ParagraphStyle(
            "SectionHeading",
            parent=styles["Heading1"],
            fontSize=14,
            textColor=colors.HexColor("#1a1a1a"),
            spaceAfter=12,
            spaceBefore=6,
        )
        story.append(Paragraph(section_title, heading_style))

        # Parse and add content
        lines = content.split("\n")
        body_style = styles["Normal"]
        bullet_style = ParagraphStyle(
            "BulletStyle",
            parent=styles["Normal"],
            leftIndent=20,
            spaceAfter=6,
        )

        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 0.1 * inch))
            elif line.startswith("- "):
                # Bullet point
                bullet_text = self._sanitize_text(line[2:])
                story.append(Paragraph(f"• {bullet_text}", bullet_style))
            elif line.endswith(":"):
                # Subheading
                subheading_style = ParagraphStyle(
                    "SubHeading",
                    parent=styles["Normal"],
                    fontSize=11,
                    textColor=colors.HexColor("#333333"),
                    spaceAfter=6,
                    spaceBefore=6,
                )
                sanitized_line = self._sanitize_text(line)
                story.append(Paragraph(f"<b>{sanitized_line}</b>", subheading_style))
            else:
                # Regular paragraph
                sanitized_line = self._sanitize_text(line)
                story.append(Paragraph(sanitized_line, body_style))

        story.append(Spacer(1, 0.2 * inch))

    def load_template(self, template_name: str) -> Optional[Document]:
        """Load a Word document template.
        
        Args:
            template_name: Name of the template file (e.g., "cam_template.docx")
            
        Returns:
            Loaded Document object or None if template not found
        """
        template_path = os.path.join(self.template_dir, template_name)
        if os.path.exists(template_path):
            try:
                return Document(template_path)
            except Exception as e:
                raise IOError(f"Failed to load template {template_name}: {str(e)}")
        return None

    def create_template(self, template_name: str) -> Document:
        """Create a new Word document template.
        
        Args:
            template_name: Name for the new template file
            
        Returns:
            New Document object
        """
        doc = Document()
        
        # Add template structure
        doc.add_heading("Credit Appraisal Memo Template", level=0)
        doc.add_paragraph("This is a template for CAM documents.")
        
        # Add placeholder sections
        sections = [
            "Executive Summary",
            "Company Overview",
            "Industry Analysis",
            "Financial Analysis",
            "Risk Assessment",
            "Five Cs Summary",
            "Final Recommendation",
            "Explainability Notes",
            "Audit Trail",
        ]
        
        for section in sections:
            doc.add_heading(section, level=1)
            doc.add_paragraph("[Content will be populated here]")
        
        # Save template
        template_path = os.path.join(self.template_dir, template_name)
        doc.save(template_path)
        
        return doc

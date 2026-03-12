"""Property-based tests for DocumentExporter class.

Feature: intelli-credit
"""

import os
import tempfile
from datetime import datetime

import pytest
from hypothesis import given, strategies as st, settings, HealthCheck

from app.models.cam import CAMDocument, AuditTrail, AuditEvent
from app.services.cam_generator.document_exporter import DocumentExporter


# Strategies for generating test data
@st.composite
def cam_documents(draw):
    """Generate valid CAM documents with sections."""
    # Use printable ASCII characters to avoid control character issues
    application_id = draw(st.text(
        alphabet=st.characters(blacklist_categories=('Cc', 'Cs')),
        min_size=1, 
        max_size=20
    ))
    company_name = draw(st.text(
        alphabet=st.characters(blacklist_categories=('Cc', 'Cs')),
        min_size=1, 
        max_size=100
    ))
    
    # Create audit trail
    audit_trail = AuditTrail(
        events=[
            AuditEvent(
                timestamp=datetime.now(),
                event_type="test",
                description="Test event",
                user="test_user",
            )
        ]
    )
    
    # Create CAM document
    cam_doc = CAMDocument(
        application_id=application_id,
        company_name=company_name,
        generated_date=datetime.now(),
        audit_trail=audit_trail,
        version=1,
    )
    
    # Add at least one section
    section_content = draw(st.text(
        alphabet=st.characters(blacklist_categories=('Cc', 'Cs')),
        min_size=10, 
        max_size=500
    ))
    cam_doc.sections["executive_summary"] = section_content
    
    return cam_doc


class TestDocumentExporterProperties:
    """Property-based tests for DocumentExporter."""

    # Feature: intelli-credit, Property 36: CAM Export Format Support
    @settings(deadline=None, suppress_health_check=[HealthCheck.too_slow])
    @given(cam_doc=cam_documents())
    def test_export_format_support_word_and_pdf(self, cam_doc):
        """For any created CAM document, the system should successfully export it 
        to both Word and PDF formats.
        
        **Validates: Requirements 14.3**
        """
        with tempfile.TemporaryDirectory() as temp_dir:
            exporter = DocumentExporter(template_dir=temp_dir)
            
            # Export to Word
            word_path = os.path.join(temp_dir, "test.docx")
            exporter.export_to_word(cam_doc, word_path)
            
            # Export to PDF
            pdf_path = os.path.join(temp_dir, "test.pdf")
            exporter.export_to_pdf(cam_doc, pdf_path)
            
            # Both files should exist
            assert os.path.exists(word_path), "Word export failed"
            assert os.path.exists(pdf_path), "PDF export failed"
            
            # Both files should have content
            assert os.path.getsize(word_path) > 0, "Word file is empty"
            assert os.path.getsize(pdf_path) > 0, "PDF file is empty"

    @settings(deadline=None, suppress_health_check=[HealthCheck.too_slow])
    @given(cam_doc=cam_documents())
    def test_export_preserves_application_id(self, cam_doc):
        """For any CAM document, the exported files should preserve the application ID."""
        with tempfile.TemporaryDirectory() as temp_dir:
            exporter = DocumentExporter(template_dir=temp_dir)
            
            # Export to Word
            word_path = os.path.join(temp_dir, "test.docx")
            exporter.export_to_word(cam_doc, word_path)
            
            # Read back and verify
            from docx import Document
            doc = Document(word_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            
            # The sanitized version should be in the document
            sanitized_id = exporter._sanitize_text(cam_doc.application_id)
            assert sanitized_id in full_text or len(sanitized_id) == 0, \
                f"Application ID {sanitized_id} not found in exported document"

    @settings(deadline=None, suppress_health_check=[HealthCheck.too_slow])
    @given(cam_doc=cam_documents())
    def test_export_preserves_company_name(self, cam_doc):
        """For any CAM document, the exported files should preserve the company name."""
        with tempfile.TemporaryDirectory() as temp_dir:
            exporter = DocumentExporter(template_dir=temp_dir)
            
            # Export to Word
            word_path = os.path.join(temp_dir, "test.docx")
            exporter.export_to_word(cam_doc, word_path)
            
            # Read back and verify
            from docx import Document
            doc = Document(word_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            
            # The sanitized version should be in the document
            sanitized_name = exporter._sanitize_text(cam_doc.company_name)
            assert sanitized_name in full_text or len(sanitized_name) == 0, \
                f"Company name {sanitized_name} not found in exported document"

    @settings(deadline=None, suppress_health_check=[HealthCheck.too_slow])
    @given(cam_doc=cam_documents())
    def test_export_preserves_section_content(self, cam_doc):
        """For any CAM document with sections, the exported files should preserve section content."""
        with tempfile.TemporaryDirectory() as temp_dir:
            exporter = DocumentExporter(template_dir=temp_dir)
            
            # Export to Word
            word_path = os.path.join(temp_dir, "test.docx")
            exporter.export_to_word(cam_doc, word_path)
            
            # Read back and verify
            from docx import Document
            doc = Document(word_path)
            full_text = "\n".join([p.text for p in doc.paragraphs])
            
            # Check that section content is preserved (at least some of it)
            for section_content in cam_doc.sections.values():
                # Sanitize the content
                sanitized_content = exporter._sanitize_text(section_content)
                # Check for key parts of the content
                lines = sanitized_content.split("\n")
                for line in lines[:3]:  # Check first few lines
                    if line.strip() and len(line.strip()) > 5:
                        # This line should appear somewhere in the document
                        assert line.strip() in full_text or \
                               any(word in full_text for word in line.strip().split()[:3]), \
                               f"Section content not preserved: {line.strip()}"

    @settings(deadline=None, suppress_health_check=[HealthCheck.too_slow])
    @given(
        application_id=st.text(
            alphabet=st.characters(blacklist_categories=('Cc', 'Cs')),
            min_size=1, 
            max_size=50
        ),
        company_name=st.text(
            alphabet=st.characters(blacklist_categories=('Cc', 'Cs')),
            min_size=1, 
            max_size=100
        ),
    )
    def test_export_handles_various_identifiers(self, application_id, company_name):
        """For any valid application ID and company name, export should succeed."""
        with tempfile.TemporaryDirectory() as temp_dir:
            exporter = DocumentExporter(template_dir=temp_dir)
            
            cam_doc = CAMDocument(
                application_id=application_id,
                company_name=company_name,
                generated_date=datetime.now(),
            )
            cam_doc.sections["executive_summary"] = "Test content"
            
            # Export to Word
            word_path = os.path.join(temp_dir, "test.docx")
            exporter.export_to_word(cam_doc, word_path)
            
            # Export to PDF
            pdf_path = os.path.join(temp_dir, "test.pdf")
            exporter.export_to_pdf(cam_doc, pdf_path)
            
            # Both should succeed
            assert os.path.exists(word_path)
            assert os.path.exists(pdf_path)

    @settings(deadline=None, suppress_health_check=[HealthCheck.too_slow])
    @given(
        num_sections=st.integers(min_value=1, max_value=9),
        section_size=st.integers(min_value=10, max_value=500),
    )
    def test_export_handles_various_section_counts(self, num_sections, section_size):
        """For any number of sections, export should succeed."""
        with tempfile.TemporaryDirectory() as temp_dir:
            exporter = DocumentExporter(template_dir=temp_dir)
            
            cam_doc = CAMDocument(
                application_id="APP-001",
                company_name="Test Company",
                generated_date=datetime.now(),
            )
            
            # Add multiple sections
            section_names = [
                "executive_summary",
                "company_overview",
                "industry_analysis",
                "financial_analysis",
                "risk_assessment",
                "five_cs_summary",
                "final_recommendation",
                "explainability_notes",
                "audit_trail",
            ]
            
            for i in range(min(num_sections, len(section_names))):
                cam_doc.sections[section_names[i]] = "X" * section_size
            
            # Export to Word
            word_path = os.path.join(temp_dir, "test.docx")
            exporter.export_to_word(cam_doc, word_path)
            
            # Export to PDF
            pdf_path = os.path.join(temp_dir, "test.pdf")
            exporter.export_to_pdf(cam_doc, pdf_path)
            
            # Both should succeed
            assert os.path.exists(word_path)
            assert os.path.exists(pdf_path)
            assert os.path.getsize(word_path) > 0
            assert os.path.getsize(pdf_path) > 0

"""Tests for DocumentExporter class."""

import os
import tempfile
from datetime import datetime
from pathlib import Path

import pytest
from docx import Document

from app.models.cam import CAMDocument, AuditTrail, AuditEvent
from app.models.credit_assessment import (
    RiskScore,
    LoanRecommendation,
    Explanation,
)
from app.services.cam_generator.document_exporter import DocumentExporter


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield tmpdir


@pytest.fixture
def sample_cam_document():
    """Create a sample CAM document for testing."""
    audit_trail = AuditTrail(
        events=[
            AuditEvent(
                timestamp=datetime.now(),
                event_type="data_ingestion",
                description="Financial data ingested",
                user="test_user",
            ),
            AuditEvent(
                timestamp=datetime.now(),
                event_type="calculation",
                description="Risk score calculated",
                user="test_user",
            ),
        ]
    )

    cam_doc = CAMDocument(
        application_id="APP-001",
        company_name="Test Company Ltd.",
        generated_date=datetime.now(),
        audit_trail=audit_trail,
        version=1,
    )

    # Add sections
    cam_doc.sections["executive_summary"] = """EXECUTIVE SUMMARY

Risk Assessment: MEDIUM
Overall Risk Score: 65.50/100

Loan Recommendation:
- Maximum Loan Amount: ₹50,00,000.00
- Recommended Interest Rate: 12.50%
- Limiting Constraint: DSCR

Key Risk Factors:
- Moderate debt levels
- Seasonal revenue fluctuations
- Competitive market conditions

Key Positive Factors:
- Strong promoter track record
- Stable cash flows
- Good collateral coverage"""

    cam_doc.sections["company_overview"] = """COMPANY OVERVIEW

Company Name: Test Company Ltd.
CIN: U12345AB2020PTC123456
GSTIN: 27AABCT1234H1Z0
Industry: Manufacturing
Incorporation Date: 15-01-2020

Promoters:
- Rajesh Kumar (Managing Director): 51.00% shareholding
- Priya Sharma (Director): 49.00% shareholding

MCA Filing Status:
- Authorized Capital: ₹1,00,00,000.00
- Paid-up Capital: ₹50,00,000.00
- Last Filing Date: 30-09-2023
- Compliance Status: Compliant"""

    cam_doc.sections["industry_analysis"] = """INDUSTRY ANALYSIS

Industry: Manufacturing
Market Size: ₹50,000 Crores
Growth Rate: 8.5% YoY

Market Sentiment: Positive
Positive News: 5
Negative News: 1

Key Risk Factors:
- Sector Risk: Moderate
- Regulatory Risk: Low
- Commodity Risk: Moderate"""

    cam_doc.sections["financial_analysis"] = """FINANCIAL ANALYSIS

Period: FY 2022-23

Income Statement:
- Revenue: ₹10,00,00,000.00
- Expenses: ₹8,50,00,000.00
- EBITDA: ₹1,50,00,000.00
- Net Profit: ₹75,00,000.00

Balance Sheet:
- Total Assets: ₹5,00,00,000.00
- Total Liabilities: ₹2,50,00,000.00
- Equity: ₹2,50,00,000.00

Cash Flow:
- Operating Cash Flow: ₹1,20,00,000.00

Key Financial Ratios:
- DSCR: 1.50
- Debt-to-Equity: 1.00
- LTV: 0.60"""

    cam_doc.sections["risk_assessment"] = """RISK ASSESSMENT

Overall Risk Score: 65.50/100
Risk Classification: MEDIUM

Five Cs Assessment:
- Character Score: Strong promoter background
- Capacity Score: Adequate repayment capacity
- Capital Score: Moderate equity strength
- Collateral Score: Good collateral coverage
- Conditions Score: Favorable market conditions

Top Risk Factors:
1. Seasonal revenue fluctuations
2. Competitive market dynamics
3. Commodity price volatility
4. Regulatory changes
5. Currency fluctuations"""

    cam_doc.sections["five_cs_summary"] = """FIVE CS SUMMARY

CHARACTER:
Summary: Strong promoter background with 15 years of industry experience
Key Factors:
  - No litigation history: 1.00
  - Good governance practices: 0.95
  - Strong credit bureau score: 0.90

CAPACITY:
Summary: Adequate repayment capacity with DSCR of 1.50
Key Factors:
  - Operating cash flow: 0.85
  - Debt service obligations: 0.80
  - Revenue stability: 0.75

CAPITAL:
Summary: Moderate equity strength with D/E ratio of 1.00
Key Factors:
  - Equity base: 0.70
  - Debt levels: 0.65
  - Net worth trend: 0.60

COLLATERAL:
Summary: Good collateral coverage with LTV of 0.60
Key Factors:
  - Asset quality: 0.80
  - Valuation accuracy: 0.75
  - Liquidity: 0.70

CONDITIONS:
Summary: Favorable market conditions with moderate sector risks
Key Factors:
  - Industry growth: 0.85
  - Regulatory environment: 0.80
  - Commodity exposure: 0.60"""

    cam_doc.sections["final_recommendation"] = """FINAL RECOMMENDATION

Risk Level: MEDIUM
Overall Risk Score: 65.50/100

Loan Recommendation:
- Maximum Loan Amount: ₹50,00,000.00
- Recommended Interest Rate: 12.50%
- Limiting Constraint: DSCR

Recommendation Rationale:
The loan recommendation is based on comprehensive analysis of the Five Cs of Credit:
1. Character - Promoter credibility and track record
2. Capacity - Ability to repay based on cash flow and DSCR
3. Capital - Equity strength and financial structure
4. Collateral - Security adequacy and LTV ratio
5. Conditions - External market and regulatory factors

Based on this analysis, the recommended loan amount and interest rate reflect the 
assessed risk level and provide appropriate risk-adjusted returns."""

    cam_doc.sections["explainability_notes"] = """EXPLAINABILITY NOTES

RISK_SCORE EXPLANATION:
Summary: Medium risk due to moderate debt levels and seasonal revenue patterns

Contributing Factors:
- Debt-to-Equity ratio: 0.3500
- Revenue stability: 0.2800
- Promoter track record: 0.2200
- Market conditions: 0.1500

Data Sources:
- Financial statements FY 2022-23
- MCA filings
- Industry reports
- News articles

Reasoning: The risk score reflects a balanced assessment of the company's financial 
position and market conditions. While the promoter has a strong track record and the 
company maintains adequate cash flows, seasonal revenue fluctuations and moderate 
debt levels present some risk."""

    cam_doc.sections["audit_trail"] = """AUDIT TRAIL

Total Events: 2

DATA_INGESTION (1 events):
- 2024-01-15 10:30:45: Financial data ingested

CALCULATION (1 events):
- 2024-01-15 10:35:20: Risk score calculated"""

    return cam_doc


@pytest.fixture
def exporter(temp_dir):
    """Create a DocumentExporter instance."""
    return DocumentExporter(template_dir=temp_dir)


class TestDocumentExporterInitialization:
    """Test DocumentExporter initialization."""

    def test_init_with_default_template_dir(self):
        """Test initialization with default template directory."""
        exporter = DocumentExporter()
        assert exporter.template_dir is not None
        assert "templates" in exporter.template_dir

    def test_init_with_custom_template_dir(self, temp_dir):
        """Test initialization with custom template directory."""
        exporter = DocumentExporter(template_dir=temp_dir)
        assert exporter.template_dir == temp_dir
        assert os.path.exists(temp_dir)

    def test_template_dir_created_if_not_exists(self, temp_dir):
        """Test that template directory is created if it doesn't exist."""
        custom_dir = os.path.join(temp_dir, "custom", "templates")
        exporter = DocumentExporter(template_dir=custom_dir)
        assert os.path.exists(custom_dir)


class TestExportToWord:
    """Test Word export functionality."""

    def test_export_to_word_creates_file(self, exporter, sample_cam_document, temp_dir):
        """Test that export_to_word creates a Word file."""
        output_path = os.path.join(temp_dir, "test_output.docx")
        exporter.export_to_word(sample_cam_document, output_path)
        assert os.path.exists(output_path)
        assert output_path.endswith(".docx")

    def test_export_to_word_file_is_valid(self, exporter, sample_cam_document, temp_dir):
        """Test that exported Word file is valid and readable."""
        output_path = os.path.join(temp_dir, "test_output.docx")
        exporter.export_to_word(sample_cam_document, output_path)
        
        # Try to open the file
        doc = Document(output_path)
        assert doc is not None
        assert len(doc.paragraphs) > 0

    def test_export_to_word_contains_company_name(self, exporter, sample_cam_document, temp_dir):
        """Test that exported Word file contains company name."""
        output_path = os.path.join(temp_dir, "test_output.docx")
        exporter.export_to_word(sample_cam_document, output_path)
        
        doc = Document(output_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert sample_cam_document.company_name in full_text

    def test_export_to_word_contains_application_id(self, exporter, sample_cam_document, temp_dir):
        """Test that exported Word file contains application ID."""
        output_path = os.path.join(temp_dir, "test_output.docx")
        exporter.export_to_word(sample_cam_document, output_path)
        
        doc = Document(output_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        assert sample_cam_document.application_id in full_text

    def test_export_to_word_contains_all_sections(self, exporter, sample_cam_document, temp_dir):
        """Test that exported Word file contains all sections."""
        output_path = os.path.join(temp_dir, "test_output.docx")
        exporter.export_to_word(sample_cam_document, output_path)
        
        doc = Document(output_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        
        # Check for section titles
        assert "Executive Summary" in full_text
        assert "Company Overview" in full_text
        assert "Industry Analysis" in full_text
        assert "Financial Analysis" in full_text
        assert "Risk Assessment" in full_text
        assert "Five Cs Summary" in full_text
        assert "Final Recommendation" in full_text
        assert "Explainability Notes" in full_text
        assert "Audit Trail" in full_text

    def test_export_to_word_with_empty_sections_raises_error(self, exporter, temp_dir):
        """Test that export_to_word raises error for empty CAM document."""
        empty_cam = CAMDocument(
            application_id="APP-002",
            company_name="Empty Company",
            generated_date=datetime.now(),
        )
        
        output_path = os.path.join(temp_dir, "empty.docx")
        with pytest.raises(ValueError, match="no sections"):
            exporter.export_to_word(empty_cam, output_path)

    def test_export_to_word_creates_parent_directories(self, exporter, sample_cam_document, temp_dir):
        """Test that export_to_word creates parent directories if needed."""
        output_path = os.path.join(temp_dir, "nested", "dir", "output.docx")
        exporter.export_to_word(sample_cam_document, output_path)
        assert os.path.exists(output_path)


class TestExportToPDF:
    """Test PDF export functionality."""

    def test_export_to_pdf_creates_file(self, exporter, sample_cam_document, temp_dir):
        """Test that export_to_pdf creates a PDF file."""
        output_path = os.path.join(temp_dir, "test_output.pdf")
        exporter.export_to_pdf(sample_cam_document, output_path)
        assert os.path.exists(output_path)
        assert output_path.endswith(".pdf")

    def test_export_to_pdf_file_has_content(self, exporter, sample_cam_document, temp_dir):
        """Test that exported PDF file has content."""
        output_path = os.path.join(temp_dir, "test_output.pdf")
        exporter.export_to_pdf(sample_cam_document, output_path)
        
        # Check file size (PDF should have some content)
        file_size = os.path.getsize(output_path)
        assert file_size > 1000  # PDF should be at least 1KB

    def test_export_to_pdf_with_empty_sections_raises_error(self, exporter, temp_dir):
        """Test that export_to_pdf raises error for empty CAM document."""
        empty_cam = CAMDocument(
            application_id="APP-003",
            company_name="Empty Company",
            generated_date=datetime.now(),
        )
        
        output_path = os.path.join(temp_dir, "empty.pdf")
        with pytest.raises(ValueError, match="no sections"):
            exporter.export_to_pdf(empty_cam, output_path)

    def test_export_to_pdf_creates_parent_directories(self, exporter, sample_cam_document, temp_dir):
        """Test that export_to_pdf creates parent directories if needed."""
        output_path = os.path.join(temp_dir, "nested", "dir", "output.pdf")
        exporter.export_to_pdf(sample_cam_document, output_path)
        assert os.path.exists(output_path)


class TestFormatSectionTitle:
    """Test section title formatting."""

    def test_format_section_title_executive_summary(self, exporter):
        """Test formatting of executive_summary."""
        result = exporter._format_section_title("executive_summary")
        assert result == "Executive Summary"

    def test_format_section_title_company_overview(self, exporter):
        """Test formatting of company_overview."""
        result = exporter._format_section_title("company_overview")
        assert result == "Company Overview"

    def test_format_section_title_five_cs_summary(self, exporter):
        """Test formatting of five_cs_summary."""
        result = exporter._format_section_title("five_cs_summary")
        assert result == "Five Cs Summary"

    def test_format_section_title_single_word(self, exporter):
        """Test formatting of single word section."""
        result = exporter._format_section_title("audit_trail")
        assert result == "Audit Trail"


class TestTemplateManagement:
    """Test template loading and creation."""

    def test_load_template_returns_none_if_not_found(self, exporter):
        """Test that load_template returns None if template doesn't exist."""
        result = exporter.load_template("nonexistent_template.docx")
        assert result is None

    def test_create_template_creates_file(self, exporter, temp_dir):
        """Test that create_template creates a template file."""
        template_name = "test_template.docx"
        exporter.create_template(template_name)
        
        template_path = os.path.join(temp_dir, template_name)
        assert os.path.exists(template_path)

    def test_create_template_file_is_valid(self, exporter, temp_dir):
        """Test that created template is a valid Word document."""
        template_name = "test_template.docx"
        exporter.create_template(template_name)
        
        template_path = os.path.join(temp_dir, template_name)
        doc = Document(template_path)
        assert doc is not None
        assert len(doc.paragraphs) > 0

    def test_load_template_after_creation(self, exporter, temp_dir):
        """Test that created template can be loaded."""
        template_name = "test_template.docx"
        exporter.create_template(template_name)
        
        loaded_doc = exporter.load_template(template_name)
        assert loaded_doc is not None
        assert hasattr(loaded_doc, 'paragraphs')


class TestExportIntegration:
    """Integration tests for export functionality."""

    def test_export_both_formats(self, exporter, sample_cam_document, temp_dir):
        """Test exporting to both Word and PDF formats."""
        word_path = os.path.join(temp_dir, "output.docx")
        pdf_path = os.path.join(temp_dir, "output.pdf")
        
        exporter.export_to_word(sample_cam_document, word_path)
        exporter.export_to_pdf(sample_cam_document, pdf_path)
        
        assert os.path.exists(word_path)
        assert os.path.exists(pdf_path)

    def test_export_preserves_document_metadata(self, exporter, sample_cam_document, temp_dir):
        """Test that export preserves document metadata."""
        output_path = os.path.join(temp_dir, "output.docx")
        exporter.export_to_word(sample_cam_document, output_path)
        
        doc = Document(output_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        
        assert sample_cam_document.application_id in full_text
        assert sample_cam_document.company_name in full_text
        assert str(sample_cam_document.version) in full_text

    def test_export_with_special_characters(self, exporter, temp_dir):
        """Test export with special characters in company name."""
        cam_doc = CAMDocument(
            application_id="APP-004",
            company_name="Test & Co. Ltd. (Pvt.)",
            generated_date=datetime.now(),
        )
        cam_doc.sections["executive_summary"] = "Test content with special chars: ₹, %, &"
        
        output_path = os.path.join(temp_dir, "special_chars.docx")
        exporter.export_to_word(cam_doc, output_path)
        
        assert os.path.exists(output_path)
        doc = Document(output_path)
        full_text = "\n".join([p.text for p in doc.paragraphs])
        # The & character may be escaped to &amp; in XML, so check for either form
        assert "Test & Co. Ltd. (Pvt.)" in full_text or "Test &amp; Co. Ltd. (Pvt.)" in full_text

    def test_export_with_large_content(self, exporter, temp_dir):
        """Test export with large content."""
        cam_doc = CAMDocument(
            application_id="APP-005",
            company_name="Large Content Company",
            generated_date=datetime.now(),
        )
        
        # Create large content
        large_content = "This is a test paragraph. " * 1000
        cam_doc.sections["executive_summary"] = large_content
        
        output_path = os.path.join(temp_dir, "large_content.docx")
        exporter.export_to_word(cam_doc, output_path)
        
        assert os.path.exists(output_path)
        doc = Document(output_path)
        assert len(doc.paragraphs) > 0
